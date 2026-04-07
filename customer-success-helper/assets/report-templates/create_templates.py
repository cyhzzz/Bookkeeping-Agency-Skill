#!/usr/bin/env python3
"""Create DOCX report templates for customer-success-helper skill."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_client_onboarding_confirmation():
    """Create client-onboarding-confirmation.docx"""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('代理记账服务合同')
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph()

    # Party A info
    p = doc.add_paragraph()
    p.add_run('甲方（客户）信息').bold = True
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ('公司名称', ''),
        ('统一社会信用代码', ''),
        ('法定代表人', ''),
        ('注册地址', ''),
        ('联系人/电话', '')
    ]
    for i, (label, value) in enumerate(data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        set_cell_shading(table.rows[i].cells[0], 'E8E8E8')

    doc.add_paragraph()

    # Party B info
    p = doc.add_paragraph()
    p.add_run('乙方（代账公司）信息').bold = True
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ('公司名称', ''),
        ('统一社会信用代码', ''),
        ('法定代表人', ''),
        ('注册地址', ''),
        ('联系人/电话', '')
    ]
    for i, (label, value) in enumerate(data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        set_cell_shading(table.rows[i].cells[0], 'E8E8E8')

    doc.add_paragraph()

    # Service content
    p = doc.add_paragraph()
    p.add_run('服务内容').bold = True
    service_items = [
        '每月记账凭证编制',
        '月度财务报表编制',
        '年度企业所得税汇算清缴',
        '年度工商年报公示',
        '税务申报（增值税、附加税、个人所得税等）',
        '财务咨询与建议'
    ]
    for item in service_items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Service period
    p = doc.add_paragraph()
    p.add_run('服务期限').bold = True
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '起始日期'
    table.rows[0].cells[1].text = ''
    table.rows[1].cells[0].text = '终止日期'
    table.rows[1].cells[1].text = ''
    set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
    set_cell_shading(table.rows[1].cells[0], 'E8E8E8')

    doc.add_paragraph()

    # Service fee
    p = doc.add_paragraph()
    p.add_run('服务费用').bold = True
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '年度服务费'
    table.rows[0].cells[1].text = '人民币：_____元/年'
    table.rows[1].cells[0].text = '付款方式'
    table.rows[1].cells[1].text = '□月付  □季付  □年付'
    set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
    set_cell_shading(table.rows[1].cells[0], 'E8E8E8')

    doc.add_paragraph()

    # Signatures
    p = doc.add_paragraph()
    p.add_run('双方签章').bold = True
    doc.add_paragraph()
    doc.add_paragraph('甲方（盖章）：__________________    日期：___________')
    doc.add_paragraph()
    doc.add_paragraph('乙方（盖章）：__________________    日期：___________')

    doc.add_page_break()

    # Appendix 1
    p = doc.add_paragraph()
    p.add_run('附件1：客户资料清单确认函').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph('感谢您选择我们的代理记账服务。为确保服务质量，请确认以下资料清单：')
    items = [
        '营业执照副本复印件',
        '公司章程',
        '法人身份证复印件',
        '银行开户许可证',
        '前年度财务报表（若有）',
        '最近一个月银行对账单',
        '本月工资表',
        '库存商品明细（如有）'
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('客户确认（签字）：___________    日期：___________')

    doc.add_page_break()

    # Appendix 2
    p = doc.add_paragraph()
    p.add_run('附件2：会计科目表确认单').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph('根据贵公司行业特点，建议使用以下会计科目表结构：')
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '科目类别'
    table.rows[0].cells[1].text = '说明'
    set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
    set_cell_shading(table.rows[0].cells[1], 'E8E8E8')
    categories = [
        ('资产类', '货币资金、应收账款、存货等'),
        ('负债类', '应付账款、短期借款、长期借款等'),
        ('所有者权益', '实收资本、资本公积、未分配利润等'),
        ('成本类', '生产成本、制造费用等'),
        ('损益类-收入', '主营业务收入、其他业务收入等'),
        ('损益类-成本', '主营业务成本、其他业务成本等'),
        ('损益类-费用', '管理费用、销售费用、财务费用等')
    ]
    for i, (cat, desc) in enumerate(categories, 1):
        table.rows[i].cells[0].text = cat
        table.rows[i].cells[1].text = desc

    doc.add_page_break()

    # Appendix 3
    p = doc.add_paragraph()
    p.add_run('附件3：税种核定确认单').bold = True
    doc.add_paragraph()
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['税种', '核定情况', '申报周期']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E8E8E8')
    tax_types = [
        ('增值税', '□一般纳税人  □小规模纳税人', '月/季'),
        ('企业所得税', '□查账征收  □核定征收', '季'),
        ('个人所得税', '代扣代缴', '月'),
        ('附加税', '随增值税', '月/季'),
        ('其他', '', '')
    ]
    for i, (tax, status, period) in enumerate(tax_types, 1):
        table.rows[i].cells[0].text = tax
        table.rows[i].cells[1].text = status
        table.rows[i].cells[2].text = period

    doc.add_page_break()

    # Appendix 4
    p = doc.add_paragraph()
    p.add_run('附件4：关键日期备忘日历').bold = True
    doc.add_paragraph()
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '日期节点'
    table.rows[0].cells[1].text = '事项说明'
    set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
    set_cell_shading(table.rows[0].cells[1], 'E8E8E8')
    key_dates = [
        ('每月15日前', '个人所得税申报截止'),
        ('每月15日前', '增值税及附加税申报截止（季报）'),
        ('每季度末15日前', '企业所得税预缴申报'),
        ('每年5月31日前', '企业所得税年度汇算清缴截止'),
        ('每年1月1日-6月30日', '工商年报公示期'),
        ('每年3月1日-6月30日', '个人所得税综合所得年度汇算'),
        ('每月结束后10日内', '银行对账单获取'),
        ('每季度结束后15日内', '社保、公积金申报')
    ]
    for i, (date, desc) in enumerate(key_dates, 1):
        table.rows[i].cells[0].text = date
        table.rows[i].cells[1].text = desc

    output_path = 'D:/project/skills/代理记账/customer-success-helper/assets/report-templates/client-onboarding-confirmation.docx'
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_customer_value_report():
    """Create customer-value-report-template.docx"""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('客户价值分析报告')
    run.bold = True
    run.font.size = Pt(26)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('XXX公司')
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph()

    # Execution summary
    p = doc.add_paragraph()
    p.add_run('执行摘要').bold = True
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    summary_data = [
        ('客户基本信息', 'XXX公司 | 行业：XX | 规模：XX人'),
        ('服务起始时间', 'XXXX年XX月'),
        ('客户分级', '□钻石客户  □金牌客户  □银牌客户')
    ]
    for i, (label, value) in enumerate(summary_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        set_cell_shading(table.rows[i].cells[0], 'E8E8E8')

    doc.add_page_break()

    # Chapter 1
    p = doc.add_paragraph()
    p.add_run('第一章 服务价值量化').bold = True

    # 1.1
    p = doc.add_paragraph()
    p.add_run('1.1 税负节省测算').bold = True
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '节省项目'
    table.rows[0].cells[1].text = '金额（元）'
    set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
    set_cell_shading(table.rows[0].cells[1], 'E8E8E8')
    data = [
        ('研发费用加计扣除', 'XX,XXX'),
        ('高新技术企业所得税优惠', 'XX,XXX'),
        ('小微企业税收优惠', 'XX,XXX')
    ]
    for i, (item, amount) in enumerate(data, 1):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = amount

    doc.add_paragraph()

    # 1.2
    p = doc.add_paragraph()
    p.add_run('1.2 避免罚款金额').bold = True
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['风险类型', '潜在罚款金额', '避免原因']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E8E8E8')
    risk_data = [
        ('逾期申报罚款', 'XX,XXX', '准时申报'),
        ('税务申报错误', 'XX,XXX', '专业审核'),
        ('资料保存不当', 'XX,XXX', '合规管理')
    ]
    for i, (risk, penalty, reason) in enumerate(risk_data, 1):
        table.rows[i].cells[0].text = risk
        table.rows[i].cells[1].text = penalty
        table.rows[i].cells[2].text = reason

    doc.add_paragraph()

    # 1.3
    p = doc.add_paragraph()
    p.add_run('1.3 申报效率提升').bold = True
    p = doc.add_paragraph('通过流程优化和专业化操作，预计为贵司节省财务工时：')
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '年度节省工时'
    table.rows[0].cells[1].text = '折算价值（元）'
    set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
    set_cell_shading(table.rows[0].cells[1], 'E8E8E8')
    eff_data = [('XX小时', 'XX,XXX元')]
    for i, (hours, value) in enumerate(eff_data, 1):
        table.rows[i].cells[0].text = hours
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 1.4
    p = doc.add_paragraph()
    p.add_run('1.4 合规风险降低').bold = True
    p = doc.add_paragraph('通过专业的风险管理和预警机制，本年度成功降低合规风险：')
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '风险类型'
    table.rows[0].cells[1].text = '处理情况'
    set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
    set_cell_shading(table.rows[0].cells[1], 'E8E8E8')
    risk_mgmt = [
        ('税务异常预警', '已及时处理'),
        ('申报数据异常', '已核实修正')
    ]
    for i, (risk, status) in enumerate(risk_mgmt, 1):
        table.rows[i].cells[0].text = risk
        table.rows[i].cells[1].text = status

    doc.add_page_break()

    # Chapter 2
    p = doc.add_paragraph()
    p.add_run('第二章 行业对标分析').bold = True

    # 2.1
    p = doc.add_paragraph()
    p.add_run('2.1 同行业税负率对比').bold = True
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['指标', '贵司', '行业平均']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E8E8E8')
    comparison = [
        ('税负率', 'XX%', 'XX%'),
        ('在行业中排名', '前XX%', '-')
    ]
    for i, (indicator, company, industry) in enumerate(comparison, 1):
        table.rows[i].cells[0].text = indicator
        table.rows[i].cells[1].text = company
        table.rows[i].cells[2].text = industry

    doc.add_paragraph()

    # 2.2
    p = doc.add_paragraph()
    p.add_run('2.2 同行业服务费用对比').bold = True
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['服务类型', '贵司费用', '市场均价']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E8E8E8')
    fee_comparison = [
        ('基础代账', 'XX,XXX元/年', 'XX,XXX元/年'),
        ('增值服务', 'XX,XXX元/年', 'XX,XXX元/年')
    ]
    for i, (service, company_fee, market_fee) in enumerate(fee_comparison, 1):
        table.rows[i].cells[0].text = service
        table.rows[i].cells[1].text = company_fee
        table.rows[i].cells[2].text = market_fee

    doc.add_paragraph()

    # 2.3
    p = doc.add_paragraph()
    p.add_run('2.3 服务满意度对比').bold = True
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '满意度维度'
    table.rows[0].cells[1].text = '评分'
    set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
    set_cell_shading(table.rows[0].cells[1], 'E8E8E8')
    satisfaction = [
        ('服务质量', 'XX/100'),
        ('行业排名', '前XX%')
    ]
    for i, (dimension, score) in enumerate(satisfaction, 1):
        table.rows[i].cells[0].text = dimension
        table.rows[i].cells[1].text = score

    doc.add_page_break()

    # Chapter 3
    p = doc.add_paragraph()
    p.add_run('第三章 客户全景图').bold = True

    # 3.1
    p = doc.add_paragraph()
    p.add_run('3.1 基本信息').bold = True
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    basic_info = [
        ('客户名称', ''),
        ('所属行业', ''),
        ('企业规模', '□小型  □中型  □大型'),
        ('客户等级', '□钻石  □金牌  □银牌'),
        ('生命周期阶段', '□初创期  □成长期  □成熟期  □转型期'),
        ('合作时长', 'X年X个月')
    ]
    for i, (label, value) in enumerate(basic_info):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        set_cell_shading(table.rows[i].cells[0], 'E8E8E8')

    doc.add_paragraph()

    # 3.2
    p = doc.add_paragraph()
    p.add_run('3.2 风险等级').bold = True
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '风险等级'
    table.rows[0].cells[1].text = '□低风险  □中风险  □高风险'
    table.rows[1].cells[0].text = '风险预警'
    table.rows[1].cells[1].text = '□无  □橙色预警  □红色预警'
    table.rows[2].cells[0].text = '风险事件（年度）'
    table.rows[2].cells[1].text = 'X次'
    for i in range(3):
        set_cell_shading(table.rows[i].cells[0], 'E8E8E8')

    doc.add_paragraph()

    # 3.3
    p = doc.add_paragraph()
    p.add_run('3.3 服务状态').bold = True
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    service_status = [
        ('合同状态', '□有效  □即将到期  □已到期'),
        ('服务套餐', ''),
        ('本月服务完成率', 'XX%'),
        ('累计服务工时', 'XX小时')
    ]
    for i, (label, value) in enumerate(service_status):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        set_cell_shading(table.rows[i].cells[0], 'E8E8E8')

    doc.add_paragraph()

    # 3.4
    p = doc.add_paragraph()
    p.add_run('3.4 关键联系人').bold = True
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    headers = ['姓名', '职位', '联系方式']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E8E8E8')
    contacts = [
        ('', '', ''),
        ('', '', '')
    ]
    for i, (name, position, contact) in enumerate(contacts, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = position
        table.rows[i].cells[2].text = contact

    doc.add_page_break()

    # Chapter 4
    p = doc.add_paragraph()
    p.add_run('第四章 未来服务建议').bold = True

    p = doc.add_paragraph()
    p.add_run('增值服务推荐').bold = True

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['推荐服务', '预计收益', '优先级', '推荐理由']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E8E8E8')
    services = [
        ('', '', '□高  □中  □低', ''),
        ('', '', '□高  □中  □低', ''),
        ('', '', '□高  □中  □低', ''),
        ('', '', '□高  □中  □低', '')
    ]
    for i, (service, benefit, priority, reason) in enumerate(services, 1):
        table.rows[i].cells[0].text = service
        table.rows[i].cells[1].text = benefit
        table.rows[i].cells[2].text = priority
        table.rows[i].cells[3].text = reason

    output_path = 'D:/project/skills/代理记账/customer-success-helper/assets/report-templates/customer-value-report-template.docx'
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_satisfaction_survey_report():
    """Create satisfaction-survey-report-template.docx"""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('满意度调研报告')
    run.bold = True
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('20XX年第X季度')
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Chapter 1
    p = doc.add_paragraph()
    p.add_run('第一章 调研概况').bold = True

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    overview = [
        ('调研时间', '20XX年XX月XX日 - XX月XX日'),
        ('调研对象', '全体服务客户（共XX家）'),
        ('回收率', 'XX%')
    ]
    for i, (label, value) in enumerate(overview):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        set_cell_shading(table.rows[i].cells[0], 'E8E8E8')

    doc.add_page_break()

    # Chapter 2
    p = doc.add_paragraph()
    p.add_run('第二章 整体满意度').bold = True

    p = doc.add_paragraph()
    p.add_run('2.1 平均分（各项维度）').bold = True
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['满意度维度', '平均分（5分制）', '趋势']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E8E8E8')
    dimensions = [
        ('服务质量', '4.XX', '↑持平 ↓'),
        ('专业度', '4.XX', '↑持平 ↓'),
        ('响应速度', '4.XX', '↑持平 ↓'),
        ('价值感知', '4.XX', '↑持平 ↓'),
        ('整体满意度', '4.XX', '↑持平 ↓')
    ]
    for i, (dim, score, trend) in enumerate(dimensions, 1):
        table.rows[i].cells[0].text = dim
        table.rows[i].cells[1].text = score
        table.rows[i].cells[2].text = trend

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('2.2 满意度分布图').bold = True
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '满意度等级'
    table.rows[0].cells[1].text = '客户占比'
    set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
    set_cell_shading(table.rows[0].cells[1], 'E8E8E8')
    distribution = [
        ('非常满意（5分）', 'XX%'),
        ('满意（4分）', 'XX%'),
        ('一般（3分）', 'XX%'),
        ('不满意（2分）', 'XX%'),
        ('非常不满意（1分）', 'XX%')
    ]
    for i, (level, pct) in enumerate(distribution, 1):
        table.rows[i].cells[0].text = level
        table.rows[i].cells[1].text = pct

    doc.add_page_break()

    # Chapter 3
    p = doc.add_paragraph()
    p.add_run('第三章 分项分析').bold = True

    dimensions_detail = [
        ('服务质量', 'XX%'),
        ('专业度', 'XX%'),
        ('响应速度', 'XX%'),
        ('价值感知', 'XX%')
    ]

    for dim, score in dimensions_detail:
        p = doc.add_paragraph()
        p.add_run(f'3.X {dim}（满意度：{score}）').bold = True
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = '评价要点'
        table.rows[0].cells[1].text = '客户反馈摘要'
        set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
        set_cell_shading(table.rows[0].cells[1], 'E8E8E8')
        table.rows[1].cells[0].text = '优点'
        table.rows[1].cells[1].text = ''
        table.rows[2].cells[0].text = '改进空间'
        table.rows[2].cells[1].text = ''
        doc.add_paragraph()

    doc.add_page_break()

    # Chapter 4
    p = doc.add_paragraph()
    p.add_run('第四章 客户反馈汇总').bold = True

    p = doc.add_paragraph()
    p.add_run('4.1 正面反馈摘录').bold = True
    for i in range(3):
        p = doc.add_paragraph(f'"{i+1}. "客户反馈内容摘录..."', style='Quote')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('4.2 负面反馈摘录').bold = True
    for i in range(3):
        p = doc.add_paragraph(f'"{i+1}. "客户反馈内容摘录..."', style='Quote')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('4.3 改进建议').bold = True
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['问题类型', '改进建议', '责任部门']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E8E8E8')
    suggestions = [
        ('', '', ''),
        ('', '', ''),
        ('', '', '')
    ]
    for i, (prob, sug, dept) in enumerate(suggestions, 1):
        table.rows[i].cells[0].text = prob
        table.rows[i].cells[1].text = sug
        table.rows[i].cells[2].text = dept

    doc.add_page_break()

    # Chapter 5
    p = doc.add_paragraph()
    p.add_run('第五章 下季度改进计划').bold = True

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['优先级', '改进事项', '负责人', '完成时间']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E8E8E8')
    plans = [
        ('□高  □中  □低', '', '', ''),
        ('□高  □中  □低', '', '', ''),
        ('□高  □中  □低', '', '', ''),
        ('□高  □中  □低', '', '', '')
    ]
    for i, (priority, item, owner, deadline) in enumerate(plans, 1):
        table.rows[i].cells[0].text = priority
        table.rows[i].cells[1].text = item
        table.rows[i].cells[2].text = owner
        table.rows[i].cells[3].text = deadline

    output_path = 'D:/project/skills/代理记账/customer-success-helper/assets/report-templates/satisfaction-survey-report-template.docx'
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_churn_risk_report():
    """Create customer-churn-risk-report-template.docx"""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('流失风险预警报告')
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph()

    # Execution summary
    p = doc.add_paragraph()
    p.add_run('执行摘要').bold = True
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    summary_data = [
        ('高风险客户数量', 'X家'),
        ('中风险客户数量', 'X家'),
        ('本月新增预警', 'X家')
    ]
    for i, (label, value) in enumerate(summary_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        set_cell_shading(table.rows[i].cells[0], 'E8E8E8')

    doc.add_page_break()

    # Chapter 1
    p = doc.add_paragraph()
    p.add_run('第一章 流失风险排名').bold = True

    p = doc.add_paragraph()
    p.add_run('1.1 高风险客户清单（红色预警）').bold = True
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    headers = ['客户名称', '风险得分', '主要风险因素', '预警时间', '负责销售']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'FFCCCC')  # Light red
    for i in range(1, 6):
        for j in range(5):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('1.2 中风险客户清单（橙色预警）').bold = True
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'FFE0CC')  # Light orange
    for i in range(1, 6):
        for j in range(5):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('1.3 风险变化趋势').bold = True
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['月份', '高风险', '中风险', '低风险']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E8E8E8')
    for i in range(1, 4):
        for j in range(4):
            table.rows[i].cells[j].text = ''

    doc.add_page_break()

    # Chapter 2
    p = doc.add_paragraph()
    p.add_run('第二章 流失原因分析').bold = True

    p = doc.add_paragraph()
    p.add_run('2.1 主动流失因素').bold = True
    factors = [
        '服务费用超出预算',
        '对服务质量不满意',
        '选择了竞争对手',
        '内部调整/裁员'
    ]
    for f in factors:
        p = doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('2.2 被动流失因素').bold = True
    factors = [
        '企业注销/破产',
        '业务转型（不再需要服务）',
        '被竞争对手收购',
        '长期不配合工作'
    ]
    for f in factors:
        p = doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('2.3 行业影响因素').bold = True
    factors = [
        '行业整体下行',
        '政策变化影响',
        '市场竞争加剧',
        '原材料价格上涨'
    ]
    for f in factors:
        p = doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # Chapter 3
    p = doc.add_paragraph()
    p.add_run('第三章 挽回建议').bold = True

    p = doc.add_paragraph()
    p.add_run('3.1 高风险客户：一对一拜访计划').bold = True
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['客户名称', '拜访计划', '负责人', '预计挽回概率']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'FFCCCC')
    for i in range(1, 4):
        for j in range(4):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('3.2 中风险客户：主动关怀计划').bold = True
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['客户名称', '关怀措施', '执行时间']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'FFE0CC')
    for i in range(1, 4):
        for j in range(3):
            table.rows[i].cells[j].text = ''

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('3.3 一般客户：标准化服务升级').bold = True
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '服务升级措施'
    table.rows[0].cells[1].text = '说明'
    set_cell_shading(table.rows[0].cells[0], 'E8E8E8')
    set_cell_shading(table.rows[0].cells[1], 'E8E8E8')
    upgrades = [
        ('定期回访', '每月主动联系一次'),
        ('增值服务赠送', '提供免费税务健康诊断')
    ]
    for i, (measure, desc) in enumerate(upgrades, 1):
        table.rows[i].cells[0].text = measure
        table.rows[i].cells[1].text = desc

    output_path = 'D:/project/skills/代理记账/customer-success-helper/assets/report-templates/customer-churn-risk-report-template.docx'
    doc.save(output_path)
    print(f"Created: {output_path}")


if __name__ == '__main__':
    create_client_onboarding_confirmation()
    create_customer_value_report()
    create_satisfaction_survey_report()
    create_churn_risk_report()
    print("\nAll templates created successfully!")
