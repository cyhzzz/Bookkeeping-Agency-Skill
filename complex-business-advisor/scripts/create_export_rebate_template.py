"""
创建出口退税报告模板
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_export_rebate_report_template():
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 封面
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('出口退税计算报告')
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('（出口货物及发票汇总 / 退税计算明细 / 风险审核报告）')
    sub_run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    cover_info = doc.add_paragraph()
    cover_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_info.add_run('\n\n\n企业名称：____________________\n\n')
    cover_info.add_run('报告日期：____________________\n\n')
    cover_info.add_run('编制人员：____________________\n\n')
    cover_info.add_run('复核人员：____________________')

    doc.add_page_break()

    # 第一章 出口货物及发票汇总
    h1 = doc.add_heading('第一章 出口货物及发票汇总', level=1)

    h2 = doc.add_heading('一、报关单汇总表', level=2)

    # 报关单汇总表
    table1 = doc.add_table(rows=6, cols=6)
    table1.style = 'Table Grid'

    # 表头
    headers1 = ['序号', '报关单号', '商品名称', '数量', '离岸价(FOB)', '目的国']
    for i, header in enumerate(headers1):
        cell = table1.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # 示例数据行
    data1 = [
        ['1', 'BJ202401001', '电子产品', '10,000', '50,000.00 USD', '美国'],
        ['2', 'BJ202401002', '纺织品', '50,000', '30,000.00 USD', '欧盟'],
        ['3', 'BJ202401003', '机械设备', '100', '80,000.00 USD', '日本'],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
    ]
    for row_idx, row_data in enumerate(data1):
        for col_idx, cell_text in enumerate(row_data):
            table1.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_sum = doc.add_heading('二、采购发票汇总表', level=2)

    # 采购发票汇总表
    table2 = doc.add_table(rows=5, cols=6)
    table2.style = 'Table Grid'

    headers2 = ['序号', '发票号', '商品名称', '价税合计', '税率', '进项税额']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    data2 = [
        ['1', 'FP202401001', '原材料A', '113,000.00', '13%', '13,000.00'],
        ['2', 'FP202401002', '原材料B', '56,500.00', '13%', '6,500.00'],
        ['3', 'FP202401003', '原材料C', '22,600.00', '13%', '2,600.00'],
        ['', '', '', '', '', ''],
    ]
    for row_idx, row_data in enumerate(data2):
        for col_idx, cell_text in enumerate(row_data):
            table2.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_page_break()

    # 第二章 退税计算明细
    h1_2 = doc.add_heading('第二章 退税计算明细', level=1)

    h2_2 = doc.add_heading('一、退税率查询表', level=2)

    table3 = doc.add_table(rows=4, cols=4)
    table3.style = 'Table Grid'

    headers3 = ['商品HS编码', '商品名称', '退税率', '适用增值税率']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    data3 = [
        ['8471', '电子产品', '13%', '13%'],
        ['6201', '纺织品', '13%', '13%'],
        ['8501', '机械设备', '13%', '13%'],
    ]
    for row_idx, row_data in enumerate(data3):
        for col_idx, cell_text in enumerate(row_data):
            table3.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_2_2 = doc.add_heading('二、汇率取值记录', level=2)
    p = doc.add_paragraph()
    p.add_run('汇率来源：').bold = True
    p.add_run('中国人民银行外汇牌价\n')
    p.add_run('申报当期汇率（中间价）：').bold = True
    p.add_run('____（填写具体汇率）\n')
    p.add_run('汇率取值日期：').bold = True
    p.add_run('____年__月__日')

    doc.add_paragraph()

    h2_2_3 = doc.add_heading('三、退税金额计算表', level=2)

    table4 = doc.add_table(rows=6, cols=4)
    table4.style = 'Table Grid'

    headers4 = ['项目', '计算公式', '金额（元）', '备注']
    for i, header in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    calc_data = [
        ['出口离岸价(FOB)', '∑各报关单FOB美元×汇率', '', ''],
        ['免抵退税额', 'FOB人民币×退税率', '', ''],
        ['不得免征抵扣税额', 'FOB人民币×(退税率-适用税率)', '', ''],
        ['应纳税额', '销项-进项-不得免征抵扣-上期留抵', '', ''],
        ['应退税额', 'MIN(应纳税额绝对值,免抵退税额)', '', ''],
    ]
    for row_idx, row_data in enumerate(calc_data):
        for col_idx, cell_text in enumerate(row_data):
            table4.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_page_break()

    # 第三章 风险审核报告
    h1_3 = doc.add_heading('第三章 风险审核报告', level=1)

    h2_3_1 = doc.add_heading('一、单证匹配检查结果', level=2)

    table5 = doc.add_table(rows=4, cols=3)
    table5.style = 'Table Grid'

    headers5 = ['检查项目', '检查结果', '是否通过']
    for i, header in enumerate(headers5):
        cell = table5.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    match_data = [
        ['报关单与发票匹配率', '____%（要求≥98%）', '□ 通过 □ 需补充'],
        ['单证齐全性', '____份/应____份', '□ 通过 □ 需补充'],
        ['商品名称规格匹配', '□ 完全匹配 □ 部分匹配', '□ 通过 □ 需调整'],
    ]
    for row_idx, row_data in enumerate(match_data):
        for col_idx, cell_text in enumerate(row_data):
            table5.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_3_2 = doc.add_heading('二、换汇成本检查', level=2)

    p2 = doc.add_paragraph()
    p2.add_run('换汇成本计算公式：').bold = True
    p2.add_run('采购不含税进价 ÷ 出口美元离岸价\n\n')
    p2.add_run('换汇成本区间判定：\n')
    p2.add_run('• <5：成本异常偏低，可能存在骗税风险\n')
    p2.add_run('• 5-8：正常区间\n')
    p2.add_run('• >8：成本偏高，出口竞争力下降\n\n')
    p2.add_run('检查结果：换汇成本 = ____（要求在5-8区间内）')

    doc.add_paragraph()

    h2_3_3 = doc.add_heading('三、异常预警清单', level=2)

    table6 = doc.add_table(rows=4, cols=4)
    table6.style = 'Table Grid'

    headers6 = ['预警级别', '预警项目', '预警详情', '处理建议']
    for i, header in enumerate(headers6):
        cell = table6.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    warning_data = [
        ['□ 高风险 □ 中风险 □ 低风险', '', '', ''],
        ['', '', '', ''],
        ['', '', '', ''],
    ]
    for row_idx, row_data in enumerate(warning_data):
        for col_idx, cell_text in enumerate(row_data):
            table6.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_page_break()

    # 第四章 申报建议
    h1_4 = doc.add_heading('第四章 申报建议', level=1)

    h2_4_1 = doc.add_heading('一、操作步骤', level=2)

    steps = [
        '1. 整理报关单、发票、装箱单、提单等单证',
        '2. 核对单证匹配率，确保≥98%',
        '3. 登录电子税务局出口退税申报模块',
        '4. 录入报关单数据及发票数据',
        '5. 确认退税计算结果',
        '6. 提交申报并等待审核'
    ]
    for step in steps:
        doc.add_paragraph(step)

    doc.add_paragraph()

    h2_4_2 = doc.add_heading('二、时间节点', level=2)

    table7 = doc.add_table(rows=4, cols=2)
    table7.style = 'Table Grid'

    time_data = [
        ['事项', '时间要求'],
        ['月度申报截止', '每月15日前（遇节假日顺延）'],
        ['申报数据确认', '申报后5个工作日内'],
        ['退税到账', '审核通过后7-15个工作日'],
    ]
    for row_idx, row_data in enumerate(time_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table7.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # 双人复核签字区
    review_section = doc.add_paragraph()
    review_section.add_run('\n\n双人复核状态：\n').bold = True
    review_section.add_run('□ AI处理完成    □ 人工确认完成\n\n')
    review_section.add_run('AI处理人：__________    日期：__________\n\n')
    review_section.add_run('人工复核人：__________    日期：__________')

    # 保存
    output_path = 'D:/project/skills/代理记账/complex-business-advisor/assets/report-templates/export-rebate-report-template.docx'
    doc.save(output_path)
    print(f'已创建：{output_path}')


if __name__ == '__main__':
    create_export_rebate_report_template()
