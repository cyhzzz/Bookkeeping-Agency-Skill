"""
创建研发费用加计扣除报告模板
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def create_rd_expense_deduction_report_template():
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 封面
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('研发费用加计扣除计算报告')
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('（研发项目概况 / 研发费用明细 / 加计扣除计算）')
    sub_run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    cover_info = doc.add_paragraph()
    cover_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_info.add_run('\n\n\n企业名称：____________________\n\n')
    cover_info.add_run('报告年度：____________________\n\n')
    cover_info.add_run('编制人员：____________________\n\n')
    cover_info.add_run('复核人员：____________________')

    doc.add_page_break()

    # 第一章 研发项目概况
    h1 = doc.add_heading('第一章 研发项目概况', level=1)

    h2 = doc.add_heading('一、研发项目列表', level=2)

    table1 = doc.add_table(rows=5, cols=5)
    table1.style = 'Table Grid'

    headers1 = ['序号', '项目编号', '项目名称', '项目阶段', '研发周期']
    for i, header in enumerate(headers1):
        cell = table1.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    project_data = [
        ['1', '', '', '□ 研发中 □ 已完成', ''],
        ['2', '', '', '□ 研发中 □ 已完成', ''],
        ['3', '', '', '□ 研发中 □ 已完成', ''],
        ['', '', '', '', ''],
    ]
    for row_idx, row_data in enumerate(project_data):
        for col_idx, cell_text in enumerate(row_data):
            table1.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_2 = doc.add_heading('二、项目所处阶段分布', level=2)

    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Table Grid'

    phase_data = [
        ['项目阶段', '项目数量', '占比'],
        ['研发中', '', ''],
        ['已完成', '', ''],
    ]
    for row_idx, row_data in enumerate(phase_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table2.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # 第二章 研发费用明细
    h1_2 = doc.add_heading('第二章 研发费用明细', level=1)

    h2_2_1 = doc.add_heading('一、人员人工费用', level=2)

    p = doc.add_paragraph()
    p.add_run('包含内容：').bold = True
    p.add_run('研发人员工资、奖金、津贴、补贴、基本养老保险、基本医疗保险、失业保险、工伤保险、生育保险和住房公积金\n\n')
    p.add_run('归集要求：需有完整的工资表、社保缴纳记录、劳动合同\n\n')

    table3 = doc.add_table(rows=4, cols=4)
    table3.style = 'Table Grid'

    headers3 = ['项目', '金额（元）', '占研发费用比例', '备注']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    data3 = [
        ['人员人工费用合计', '', '', ''],
        ['其中：工资薪金', '', '', ''],
        ['社保及公积金', '', '', ''],
    ]
    for row_idx, row_data in enumerate(data3):
        for col_idx, cell_text in enumerate(row_data):
            table3.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_2_2 = doc.add_heading('二、直接投入费用', level=2)

    p2 = doc.add_paragraph()
    p2.add_run('包含内容：').bold = True
    p2.add_run('材料、燃料和动力费用、模具开发费、样品样机购置费、检验费、仪器设备运行维护费等\n\n')

    table4 = doc.add_table(rows=4, cols=4)
    table4.style = 'Table Grid'

    for i, header in enumerate(headers3):
        cell = table4.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    data4 = [
        ['直接投入费用合计', '', '', ''],
        ['其中：材料费', '', '', ''],
        ['水电费', '', '', ''],
    ]
    for row_idx, row_data in enumerate(data4):
        for col_idx, cell_text in enumerate(row_data):
            table4.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_2_3 = doc.add_heading('三、折旧费用', level=2)

    table5 = doc.add_table(rows=3, cols=4)
    table5.style = 'Table Grid'

    for i, header in enumerate(headers3):
        cell = table5.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    data5 = [
        ['折旧费用合计', '', '', ''],
        ['研发设备折旧', '', '', ''],
    ]
    for row_idx, row_data in enumerate(data5):
        for col_idx, cell_text in enumerate(row_data):
            table5.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_2_4 = doc.add_heading('四、无形资产摊销', level=2)

    table6 = doc.add_table(rows=3, cols=4)
    table6.style = 'Table Grid'

    for i, header in enumerate(headers3):
        cell = table6.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    data6 = [
        ['无形资产摊销合计', '', '', ''],
        ['软件摊销', '', '', ''],
    ]
    for row_idx, row_data in enumerate(data6):
        for col_idx, cell_text in enumerate(row_data):
            table6.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_2_5 = doc.add_heading('五、其他费用', level=2)

    p5 = doc.add_paragraph()
    p5.add_run('包含内容：').bold = True
    p5.add_run('技术图书资料费、资料翻译费、专家咨询费、研发保险费、知识产权申请费等\n\n')
    p5.add_run('限价要求：').bold = True
    p5.add_run('不得超过研发费用总额的10%\n\n')

    table7 = doc.add_table(rows=3, cols=4)
    table7.style = 'Table Grid'

    for i, header in enumerate(headers3):
        cell = table7.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    data7 = [
        ['其他费用合计', '', '', ''],
        ['其他费用占比', '', '（要求≤10%）', ''],
    ]
    for row_idx, row_data in enumerate(data7):
        for col_idx, cell_text in enumerate(row_data):
            table7.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_page_break()

    # 第三章 加计扣除计算
    h1_3 = doc.add_heading('第三章 加计扣除计算', level=1)

    h2_3_1 = doc.add_heading('一、研发费用辅助账汇总', level=2)

    table8 = doc.add_table(rows=8, cols=3)
    table8.style = 'Table Grid'

    headers8 = ['费用类别', '金额（元）', '占总费用比例']
    for i, header in enumerate(headers8):
        cell = table8.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    summary_data = [
        ['人员人工', '', ''],
        ['直接投入', '', ''],
        ['折旧费用', '', ''],
        ['无形资产摊销', '', ''],
        ['其他费用', '', ''],
        ['研发费用合计', '', '100%'],
        ['其他费用占比', '', '（需≤10%）'],
    ]
    for row_idx, row_data in enumerate(summary_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table8.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            if row_idx == 5:  # 合计行
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    h2_3_2 = doc.add_heading('二、可加计扣除金额计算', level=2)

    table9 = doc.add_table(rows=5, cols=2)
    table9.style = 'Table Grid'

    calc_data = [
        ['项目', '金额/比例'],
        ['研发费用合计（元）', ''],
        ['加计扣除比例', '100%（2023年起）'],
        ['可加计扣除金额（元）', ''],
        ['适用企业所得税税率', '□ 25% □ 15%（高新企业）'],
    ]
    for row_idx, row_data in enumerate(calc_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table9.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    h2_3_3 = doc.add_heading('三、节约税款测算', level=2)

    table10 = doc.add_table(rows=3, cols=2)
    table10.style = 'Table Grid'

    saving_data = [
        ['测算场景', '节约税款（元）'],
        ['法定税率25%下', ''],
        ['高新企业税率15%下', ''],
    ]
    for row_idx, row_data in enumerate(saving_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table10.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    h2_3_4 = doc.add_heading('四、优惠备案建议', level=2)

    p6 = doc.add_paragraph()
    p6.add_run('备案方式：').bold = True
    p6.add_run('□ 季度预缴申报时享受  □ 年度汇算清缴时填报\n\n')
    p6.add_run('申报时间节点：\n')
    p6.add_run('• 季度预缴：每季度结束后15日内\n')
    p6.add_run('• 年度汇算：次年5月31日前\n\n')
    p6.add_run('适用政策依据：').bold = True
    p6.add_run('\n《财政部 税务总局关于进一步完善研发费用税前加计扣除政策的公告》（2023年第7号）')

    doc.add_page_break()

    # 第四章 政策依据
    h1_4 = doc.add_heading('第四章 政策依据', level=1)

    policies = [
        '1. 《财政部 税务总局关于进一步完善研发费用税前加计扣除政策的公告》（2023年第7号）',
        '2. 《国家税务总局关于进一步落实研发费用加计扣除政策有关问题的公告》（2021年第28号）',
        '3. 《国家税务总局关于发布<企业所得税优惠政策事项办理办法>的公告》（2018年第23号）',
        '4. 《财政部 国家税务总局 科技部关于完善研究开发费用税前加计扣除政策的通知》（2015年第119号）',
    ]

    for policy in policies:
        doc.add_paragraph(policy)

    doc.add_paragraph()

    # 备查资料清单
    h2_4_1 = doc.add_heading('备查资料清单', level=2)

    docs_list = [
        '□ 研发项目立项文件',
        '□ 研发人员名单及工资表',
        '□ 研发费用辅助账',
        '□ 研发设备折旧计算表',
        '□ 无形资产摊销计算表',
        '□ 委托研发合同（如有）',
        '□ 会计凭证及原始凭证',
    ]

    for item in docs_list:
        doc.add_paragraph(item)

    doc.add_paragraph()

    # 双人复核签字区
    review_section = doc.add_paragraph()
    review_section.add_run('\n\n双人复核状态：\n').bold = True
    review_section.add_run('□ AI处理完成    □ 人工确认完成\n\n')
    review_section.add_run('AI处理人：__________    日期：__________\n\n')
    review_section.add_run('人工复核人：__________    日期：__________')

    # 保存
    output_path = 'D:/project/skills/代理记账/complex-business-advisor/assets/report-templates/rd-expense-deduction-report-template.docx'
    doc.save(output_path)
    print(f'已创建：{output_path}')


if __name__ == '__main__':
    create_rd_expense_deduction_report_template()
