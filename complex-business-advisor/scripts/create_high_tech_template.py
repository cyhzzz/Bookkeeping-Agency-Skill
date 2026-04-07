"""
创建高新企业认定报告模板
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def create_high_tech_certification_report_template():
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 封面
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('高新企业认定自评报告')
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('（企业基本情况 / 核心指标 / 自评打分 / 改进建议）')
    sub_run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    cover_info = doc.add_paragraph()
    cover_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_info.add_run('\n\n\n企业名称：____________________\n\n')
    cover_info.add_run('报告年度：____________________\n\n')
    cover_info.add_run('编制人员：____________________\n\n')
    cover_info.add_run('复核人员：____________________')

    doc.add_page_break()

    # 第一章 企业基本情况
    h1 = doc.add_heading('第一章 企业基本情况', level=1)

    table1 = doc.add_table(rows=7, cols=2)
    table1.style = 'Table Grid'

    basic_data = [
        ['企业名称', ''],
        ['统一社会信用代码', ''],
        ['注册成立时间', ''],
        ['职工总数（人）', ''],
        ['其中：科技人员数（人）', ''],
        ['上年度总收入（万元）', ''],
        ['其中：高新产品收入（万元）', ''],
    ]
    for row_idx, row_data in enumerate(basic_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table1.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # 第二章 核心指标达标情况
    h1_2 = doc.add_heading('第二章 核心指标达标情况', level=1)

    h2_2_1 = doc.add_heading('一、高新产品收入占比', level=2)

    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Table Grid'

    headers2 = ['指标', '数值', '达标要求']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    income_data = [
        ['高新产品收入（万元）', '', ''],
        ['企业总收入（万元）', '', ''],
        ['高新收入占比', '', '≥60%'],
    ]
    for row_idx, row_data in enumerate(income_data):
        for col_idx, cell_text in enumerate(row_data):
            table2.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_2_2 = doc.add_heading('二、研发费用占比', level=2)

    table3 = doc.add_table(rows=5, cols=3)
    table3.style = 'Table Grid'

    for i, header in enumerate(headers2):
        cell = table3.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    rd_data = [
        ['近三年研发费用合计（万元）', '', ''],
        ['近一年销售收入（万元）', '', ''],
        ['研发费用占比', '', ''],
        ['销售收入所属档位', '', '对应要求'],
    ]
    for row_idx, row_data in enumerate(rd_data):
        for col_idx, cell_text in enumerate(row_data):
            table3.rows[row_idx + 1].cells[col_idx].text = cell_text

    p = doc.add_paragraph()
    p.add_run('\n研发费用占比标准：\n').bold = True
    p.add_run('• 销售收入 < 5000万元：≥5%\n')
    p.add_run('• 销售收入 5000万-2亿元：≥4%\n')
    p.add_run('• 销售收入 > 2亿元：≥3%')

    doc.add_paragraph()

    h2_2_3 = doc.add_heading('三、科技人员占比', level=2)

    table4 = doc.add_table(rows=4, cols=3)
    table4.style = 'Table Grid'

    for i, header in enumerate(headers2):
        cell = table4.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    personnel_data = [
        ['科技人员数（人）', '', ''],
        ['职工总数（人）', '', ''],
        ['科技人员占比', '', '≥10%'],
    ]
    for row_idx, row_data in enumerate(personnel_data):
        for col_idx, cell_text in enumerate(row_data):
            table4.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_page_break()

    # 第三章 自评打分
    h1_3 = doc.add_heading('第三章 自评打分', level=1)

    h2_3_1 = doc.add_heading('一、六项核心指标达标情况', level=2)

    table5 = doc.add_table(rows=7, cols=4)
    table5.style = 'Table Grid'

    headers5 = ['指标', '标准要求', '实际值', '是否达标']
    for i, header in enumerate(headers5):
        cell = table5.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    core_data = [
        ['1.成立时间', '满1年', '', '□ 是 □ 否'],
        ['2.高新收入占比', '≥60%', '', '□ 是 □ 否'],
        ['3.研发费用占比', '见附注', '', '□ 是 □ 否'],
        ['4.科技人员占比', '≥10%', '', '□ 是 □ 否'],
        ['5.企业所得税率', '15%', '', '□ 是 □ 否'],
        ['6.研发组织管理', '评分制', '', '□ 是 □ 否'],
    ]
    for row_idx, row_data in enumerate(core_data):
        for col_idx, cell_text in enumerate(row_data):
            table5.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_3_2 = doc.add_heading('二、评审指标自评', level=2)

    table6 = doc.add_table(rows=6, cols=5)
    table6.style = 'Table Grid'

    headers6 = ['评分指标', '满分', '自评得分', '得分依据', '备注']
    for i, header in enumerate(headers6):
        cell = table6.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    scoring_data = [
        ['1.知识产权', '30', '', '', ''],
        ['2.科技成果转化能力', '30', '', '', ''],
        ['3.研发组织管理水平', '20', '', '', ''],
        ['4.企业成长性', '20', '', '', ''],
        ['合计', '100', '', '要求≥71分', ''],
    ]
    for row_idx, row_data in enumerate(scoring_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table6.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            if row_idx == 4:  # 合计行
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    h2_3_3 = doc.add_heading('三、知识产权评分明细', level=2)

    table7 = doc.add_table(rows=5, cols=4)
    table7.style = 'Table Grid'

    headers7 = ['评分项', '分值', '自评得分', '说明']
    for i, header in enumerate(headers7):
        cell = table7.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    ip_data = [
        ['技术的先进程度', '≤8', '', '国际领先(8)/国际先进(6)/国内领先(4)/国内先进(2)'],
        ['对主要产品的核心支持作用', '≤8', '', '强(8)/较强(6)/一般(4)/弱(2)/无(0)'],
        ['知识产权数量', '≤8', '', 'Ⅰ类发明专利(8)/Ⅱ类5项以上(6)/3-4项(4)/1-3项(2)'],
        ['知识产权获取方式', '≤6', '', '自主研发(6)/受让并购(4)/独占许可(3)/其他(0)'],
    ]
    for row_idx, row_data in enumerate(ip_data):
        for col_idx, cell_text in enumerate(row_data):
            table7.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_3_4 = doc.add_heading('四、企业成长性评分明细', level=2)

    table8 = doc.add_table(rows=3, cols=4)
    table8.style = 'Table Grid'

    for i, header in enumerate(headers7[:3]):
        cell = table8.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    table8.rows[0].cells[3].text = '说明'
    table8.rows[0].cells[3].paragraphs[0].runs[0].bold = True

    growth_data = [
        ['净资产增长率', '≤10', '', '≥35%(9-10分)/≥25%(7-8分)/≥15%(5-6分)/≥5%(3-4分)/>0%(1-2分)'],
        ['销售收入增长率', '≤10', '', '同上'],
    ]
    for row_idx, row_data in enumerate(growth_data):
        for col_idx, cell_text in enumerate(row_data):
            table8.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_page_break()

    # 第四章 不达标项改进建议
    h1_4 = doc.add_heading('第四章 不达标项改进建议', level=1)

    h2_4_1 = doc.add_heading('一、综合评价', level=2)

    table9 = doc.add_table(rows=4, cols=2)
    table9.style = 'Table Grid'

    eval_data = [
        ['评价项目', '评价结果'],
        ['六项核心指标', '□ 全部通过  □ 部分通过（____/6）'],
        ['评审得分', '____分（要求≥71分）'],
        ['综合判定', '□ 通过  □ 待改进  □ 不通过'],
    ]
    for row_idx, row_data in enumerate(eval_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table9.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    h2_4_2 = doc.add_heading('二、不达标项改进建议', level=2)

    table10 = doc.add_table(rows=6, cols=5)
    table10.style = 'Table Grid'

    improve_headers = ['序号', '不达标项', '当前状态', '目标要求', '改进建议']
    for i, header in enumerate(improve_headers):
        cell = table10.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    improve_data = [
        ['1', '', '', '', ''],
        ['2', '', '', '', ''],
        ['3', '', '', '', ''],
        ['4', '', '', '', ''],
        ['5', '', '', '', ''],
    ]
    for row_idx, row_data in enumerate(improve_data):
        for col_idx, cell_text in enumerate(row_data):
            table10.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_paragraph()

    h2_4_3 = doc.add_heading('三、改进时间规划', level=2)

    table11 = doc.add_table(rows=5, cols=4)
    table11.style = 'Table Grid'

    plan_headers = ['改进项', '改进措施', '计划完成时间', '责任人']
    for i, header in enumerate(plan_headers):
        cell = table11.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    plan_data = [
        ['', '', '', ''],
        ['', '', '', ''],
        ['', '', '', ''],
        ['', '', '', ''],
    ]
    for row_idx, row_data in enumerate(plan_data):
        for col_idx, cell_text in enumerate(row_data):
            table11.rows[row_idx + 1].cells[col_idx].text = cell_text

    doc.add_page_break()

    # 附录
    h1_5 = doc.add_heading('附录：高新企业认定相关法规', level=1)

    regulations = [
        '1. 《高新技术企业认定管理办法》（国科发火〔2016〕32号）',
        '2. 《高新技术企业认定管理工作指引》（国科发火〔2016〕195号）',
        '3. 《国家重点支持的高新技术领域》（2016年版）',
        '4. 《企业会计准则第6号——无形资产》',
        '5. 《企业会计准则第14号——收入》',
    ]

    for reg in regulations:
        doc.add_paragraph(reg)

    doc.add_paragraph()

    # 认定流程
    h2 = doc.add_heading('高新企业认定流程', level=2)

    p = doc.add_paragraph()
    p.add_run('企业自我评价 → 注册登记 → 提交材料 → 专家评审 → 认定报备 → 公示公告 → 颁发证书\n\n')
    p.add_run('评审周期：约2-3个月\n')
    p.add_run('有效期：三年（到期需重新认定）')

    doc.add_paragraph()

    # 双人复核签字区
    review_section = doc.add_paragraph()
    review_section.add_run('\n\n双人复核状态：\n').bold = True
    review_section.add_run('□ AI处理完成    □ 人工确认完成\n\n')
    review_section.add_run('AI处理人：__________    日期：__________\n\n')
    review_section.add_run('人工复核人：__________    日期：__________')

    # 保存
    output_path = 'D:/project/skills/代理记账/complex-business-advisor/assets/report-templates/high-tech-certification-report-template.docx'
    doc.save(output_path)
    print(f'已创建：{output_path}')


if __name__ == '__main__':
    create_high_tech_certification_report_template()
