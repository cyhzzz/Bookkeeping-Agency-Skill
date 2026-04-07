"""
金蝶KIS标准导出格式解析器
支持版本: KIS14.0, KIS15.0, KIS16.0
导出格式: .docx (Word文档格式)
"""

import docx
from pathlib import Path
from typing import List, Tuple, Optional
from datetime import datetime
import re

from .base_parser import (
    BaseParser,
    ParsedFinancialData,
    ParsedVoucher,
    ParsedTrialBalance,
    VoucherEntry
)


class KingdeeParser(BaseParser):
    """金蝶KIS解析器"""

    @property
    def software_name(self) -> str:
        return "金蝶KIS"

    @property
    def supported_formats(self) -> List[str]:
        return [".docx"]

    def can_parse(self, file_path: str) -> bool:
        """判断是否能解析此文件"""
        path = Path(file_path)
        if path.suffix.lower() != '.docx':
            return False

        try:
            doc = docx.Document(file_path)
            text = '\n'.join([p.text for p in doc.paragraphs])
            return '金蝶' in text or 'KIS' in text
        except Exception:
            return False

    def parse(self, file_path: str) -> ParsedFinancialData:
        """解析金蝶导出的Word文件"""
        doc = docx.Document(file_path)

        # 1. 提取表头信息
        company_name = self._extract_company_name(doc)
        period = self._extract_period(doc)

        # 2. 解析凭证数据
        vouchers = self._parse_vouchers(doc, file_path)

        # 3. 解析科目余额表（如有）
        trial_balance = self._parse_trial_balance(doc)

        return ParsedFinancialData(
            parser_name=self.software_name,
            company_name=company_name,
            period=period,
            vouchers=vouchers,
            trial_balance=trial_balance,
            raw_data={'source_file': file_path},
            confidence=0.9
        )

    def _extract_company_name(self, doc) -> str:
        """提取公司名称"""
        for para in doc.paragraphs[:10]:
            text = para.text.strip()
            if text and len(text) < 50:
                # 跳过标题行
                if any(k in text for k in ['凭证', '科目余额', '日记账', '金蝶', 'KIS']):
                    continue
                return text
        return "未知公司"

    def _extract_period(self, doc) -> str:
        """提取会计期间"""
        for para in doc.paragraphs[:20]:
            text = para.text.strip()
            # 匹配日期格式
            match = re.search(r'(\d{4})[年-](\d{1,2})[月-]?', text)
            if match:
                return f"{match.group(1)}-{int(match.group(2)):02d}"
        return datetime.now().strftime("%Y-%m")

    def _parse_vouchers(self, doc, source_file: str) -> List[ParsedVoucher]:
        """解析凭证数据"""
        vouchers = []
        tables = doc.tables

        voucher_id = 1
        for table in tables:
            parsed_vouchers = self._parse_single_table(table, source_file, voucher_id)
            vouchers.extend(parsed_vouchers)
            voucher_id += len(parsed_vouchers)

        return vouchers

    def _parse_single_table(self, table, source_file: str, start_id: int) -> List[ParsedVoucher]:
        """解析单个表格中的凭证"""
        if not table.rows:
            return []

        vouchers = []

        # 检查是否是凭证表格
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if not self._is_voucher_header(header):
            return []

        current_voucher = None
        current_entries = []
        voucher_no = ""
        voucher_date = ""
        voucher_summary = ""

        row_count = len(table.rows)
        for i in range(1, min(row_count, 500)):  # 最多处理500行
            row = table.rows[i]
            cells = [cell.text.strip() for cell in row.cells]

            # 跳过空行
            if not any(cells):
                continue

            # 判断行类型
            row_type = self._classify_row(cells, header)

            if row_type == 'header':
                # 保存上一个凭证
                if current_voucher:
                    current_voucher.entries = current_entries
                    vouchers.append(current_voucher)

                voucher_no = cells[0] if len(cells) > 0 else f"记-{start_id}"
                voucher_date = cells[1] if len(cells) > 1 else ""
                voucher_summary = cells[2] if len(cells) > 2 else ""
                current_entries = []
                current_voucher = ParsedVoucher(
                    voucher_id=f"V{voucher_no}",
                    date=voucher_date,
                    voucher_no=voucher_no,
                    summary=voucher_summary,
                    source_file=source_file
                )
                start_id += 1

            elif row_type == 'entry' and current_voucher:
                # 解析分录
                entry = self._parse_entry_row(cells, header)
                if entry:
                    current_entries.append(entry)

        # 保存最后一个凭证
        if current_voucher:
            current_voucher.entries = current_entries
            vouchers.append(current_voucher)

        return vouchers

    def _is_voucher_header(self, header: List[str]) -> bool:
        """判断是否为凭证表头"""
        header_text = ' '.join(header).lower()
        keywords = ['日期', '凭证号', '摘要', '科目', '借方', '贷方', '金额']
        match_count = sum(1 for k in keywords if k in header_text)
        return match_count >= 3

    def _classify_row(self, cells: List[str], header: List[str]) -> str:
        """分类表格行"""
        if not cells[0]:
            return 'empty'

        first_cell = cells[0].strip()

        # 检查是否是日期格式
        if re.match(r'\d{4}[-/]\d{1,2}', first_cell):
            return 'header'

        # 检查是否是数字（金额）
        try:
            float(cells[-1].replace(',', '').replace(' ', ''))
            return 'entry'
        except (ValueError, IndexError):
            return 'empty'

    def _parse_entry_row(self, cells: List[str], header: List[str]) -> Optional[VoucherEntry]:
        """解析分录行"""
        try:
            # 建立列索引映射
            col_map = {h: i for i, h in enumerate(header)}

            account_code = ""
            account_name = ""
            debit = 0.0
            credit = 0.0

            # 查找科目列
            for i, cell in enumerate(cells):
                if i < len(header):
                    h = header[i]
                    if '科目' in h and '编码' in h:
                        account_code = cell
                    elif '科目' in h and '名称' in h:
                        account_name = cell
                    elif '借方' in h:
                        try:
                            debit = float(cell.replace(',', '').replace(' ', '') or 0)
                        except ValueError:
                            debit = 0
                    elif '贷方' in h:
                        try:
                            credit = float(cell.replace(',', '').replace(' ', '') or 0)
                        except ValueError:
                            credit = 0

            if account_name:
                return VoucherEntry(
                    account_code=account_code,
                    account_name=account_name,
                    debit_amount=debit,
                    credit_amount=credit
                )
        except Exception:
            pass
        return None

    def _parse_trial_balance(self, doc) -> List[ParsedTrialBalance]:
        """解析科目余额表"""
        # 简化实现，后续可扩展
        return []
