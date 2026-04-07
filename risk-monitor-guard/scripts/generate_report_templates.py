"""
生成报告模板DOCX文件
使用python-docx库创建专业的报告模板
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_tax_risk_alert_report():
    """创建税务风险预警报告模板"""
    doc = Document()

    # 设置文档默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ============= 封面 =============
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('税务风险预警报告')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('（XXXX年XX月）')
    run.font.size = Pt(16)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    doc.add_paragraph()

    # 报告信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run('报告编号：TAX-RISK-XXXX-XXXX\n')
    run.font.size = Pt(12)
    run = info_para.add_run(f'编制日期：XXXX年XX月XX日\n')
    run.font.size = Pt(12)
    run = info_para.add_run('编制部门：风险监控部')
    run.font.size = Pt(12)

    doc.add_page_break()

    # ============= 执行摘要 =============
    heading = doc.add_heading('执行摘要', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 风险概览表格
    p = doc.add_paragraph()
    run = p.add_run('一、风险概览（红/橙/黄/绿客户分布）')
    run.font.bold = True
    run.font.size = Pt(12)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ['风险等级', '客户数量', '占比', '环比变化']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    data = [
        ['红色（高风险）', 'X', 'X%', '↑/↓X%'],
        ['橙色（中风险）', 'X', 'X%', '↑/↓X%'],
        ['黄色（低风险）', 'X', 'X%', '↑/↓X%'],
        ['绿色（正常）', 'X', 'X%', '↑/↓X%'],
    ]
    colors = ['FFE0E0', 'FFE8CC', 'FFFFCC', 'E0FFE0']
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, colors[row_idx])

    doc.add_paragraph()

    # 本月新增重大风险
    p = doc.add_paragraph()
    run = p.add_run('二、本月新增重大风险')
    run.font.bold = True
    run.font.size = Pt(12)

    risk_table = doc.add_table(rows=4, cols=4)
    risk_table.style = 'Table Grid'
    headers = ['客户名称', '风险类型', '风险等级', '处理状态']
    for i, header in enumerate(headers):
        cell = risk_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2F5496')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 4):
        for j in range(4):
            cell = risk_table.rows[i].cells[j]
            cell.text = '示例内容' if j == 0 else ''
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============= 第一章 =============
    heading = doc.add_heading('第一章  税务风险预警报告', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 1.1 税负率异常
    heading = doc.add_heading('1.1 税负率异常客户清单', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('本节列出税负率偏离行业均值超过30%的客户清单，详细分析其税负率异常的原因及潜在风险。')

    table = doc.add_table(rows=6, cols=6)
    table.style = 'Table Grid'
    headers = ['序号', '客户名称', '行业', '实际税负率', '行业均值', '偏离度']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 6):
        for j, text in enumerate(['1', '客户A', '商贸', '1.2%', '3.0%', '-60%']):
            cell = table.rows[i].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1.2 发票合规性预警
    heading = doc.add_heading('1.2 发票合规性预警', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('本节列出发票使用存在异常的的客户，包括作废发票占比过高、红字发票频繁、异常金额发票等情况。')

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    headers = ['序号', '客户名称', '异常类型', '具体情况', '风险等级']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 5):
        for j, text in enumerate(['1', '客户B', '作废发票占比高', '作废率8.5%', '橙色']):
            cell = table.rows[i].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1.3 申报及时性预警
    heading = doc.add_heading('1.3 申报及时性预警', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('本节列出申报存在逾期的的客户，包括增值税、企业所得税、个人所得税等税种的逾期情况。')

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    headers = ['序号', '客户名称', '逾期税种', '逾期次数', '连续逾期月数']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 5):
        for j, text in enumerate(['1', '客户C', '增值税', '2次', '1个月']):
            cell = table.rows[i].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1.4 优惠备案预警
    heading = doc.add_heading('1.4 优惠备案预警', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('本节列出优惠备案存在问题的客户，包括备案过期、备案与实际不符等情况。')

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    headers = ['序号', '客户名称', '优惠类型', '问题描述', '到期日期']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 5):
        for j, text in enumerate(['1', '客户D', '高新企业', '证书即将到期', 'XXXX-XX-XX']):
            cell = table.rows[i].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============= 第二章 =============
    heading = doc.add_heading('第二章  风险趋势分析', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 2.1 行业税负率走势
    heading = doc.add_heading('2.1 行业税负率走势', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('本节分析主要行业的税负率走势变化，帮助识别系统性风险。')

    table = doc.add_table(rows=6, cols=6)
    table.style = 'Table Grid'
    headers = ['行业', '上月均值', '本月均值', '变化幅度', '预警家数', '趋势']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 6):
        for j, text in enumerate(['商贸', '2.0%', '1.8%', '-10%', '3', '↓']):
            cell = table.rows[i].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 2.2 风险转化追踪
    heading = doc.add_heading('2.2 风险转化追踪', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('本节追踪客户风险等级的变化情况，记录风险升降级情况。')

    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    headers = ['客户名称', '上月等级', '本月等级', '变化原因', '变化日期']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 6):
        for j, text in enumerate(['客户E', '黄色', '橙色', '税负率持续走低', 'XXXX-XX-XX']):
            cell = table.rows[i].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============= 第三章 =============
    heading = doc.add_heading('第三章  处理建议', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('本章针对各类风险提供具体的处理建议和操作方案。')

    # 按客户罗列
    heading = doc.add_heading('3.1 红色风险客户处理方案', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 3):
        p = doc.add_paragraph()
        run = p.add_run(f'客户X - 风险问题及处理方案')
        run.font.bold = True
        run.font.size = Pt(12)

        p = doc.add_paragraph()
        p.add_run('风险描述：')
        p.add_run('[详细描述风险内容]')

        p = doc.add_paragraph()
        p.add_run('依据条款：')
        p.add_run('[引用相关法规依据]')

        p = doc.add_paragraph()
        p.add_run('处理建议：')
        p.add_run('[具体的操作指导]')

        p = doc.add_paragraph()
        p.add_run('处理时限：')
        p.add_run('[立即处理/24小时内/本期结束前]')

        doc.add_paragraph()

    heading = doc.add_heading('3.2 橙色风险客户处理方案', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 3):
        p = doc.add_paragraph()
        run = p.add_run(f'客户Y - 风险问题及处理方案')
        run.font.bold = True
        run.font.size = Pt(12)

        p = doc.add_paragraph()
        p.add_run('风险描述：')
        p.add_run('[详细描述风险内容]')

        p = doc.add_paragraph()
        p.add_run('处理建议：')
        p.add_run('[具体的操作指导]')

        doc.add_paragraph()

    # 保存
    output_path = 'D:/project/skills/代理记账/risk-monitor-guard/assets/report-templates/tax-risk-alert-report-template.docx'
    doc.save(output_path)
    print(f'已生成: {output_path}')


def create_account_health_report():
    """创建账务健康度诊断报告模板"""
    doc = Document()

    # 设置文档默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ============= 封面 =============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('账务健康度诊断报告')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('（XXXX年XX月）')
    run.font.size = Pt(16)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run('报告编号：ACCT-HEALTH-XXXX-XXXX\n')
    run.font.size = Pt(12)
    run = info_para.add_run(f'编制日期：XXXX年XX月XX日\n')
    run.font.size = Pt(12)
    run = info_para.add_run('编制部门：账务审核部')
    run.font.size = Pt(12)

    doc.add_page_break()

    # ============= 执行摘要 =============
    heading = doc.add_heading('执行摘要', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 健康度评分分布
    p = doc.add_paragraph()
    run = p.add_run('一、健康度评分分布')
    run.font.bold = True
    run.font.size = Pt(12)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['评分区间', '客户数量', '占比']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ['90-100分（优秀）', 'X', 'X%'],
        ['75-89分（良好）', 'X', 'X%'],
        ['60-74分（合格）', 'X', 'X%'],
        ['60分以下（不合格）', 'X', 'X%'],
    ]
    colors = ['E0FFE0', 'FFFFCC', 'FFE8CC', 'FFE0E0']
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, colors[row_idx])

    doc.add_paragraph()

    # 本月重点关注客户
    p = doc.add_paragraph()
    run = p.add_run('二、本月重点关注客户')
    run.font.bold = True
    run.font.size = Pt(12)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['客户名称', '健康度评分', '主要问题', '问题类型']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2F5496')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 4):
        for j, text in enumerate(['客户A', '45分', '凭证缺失附件严重', '凭证规范性']):
            cell = table.rows[i].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============= 第一章 =============
    heading = doc.add_heading('第一章  账务健康度诊断', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 1.1 凭证规范性检查
    heading = doc.add_heading('1.1 凭证规范性检查', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('检查项目包括凭证要素完整性、编号连续性、签字完备性等。')

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['检查项目', '合格标准', '实际结果', '是否合格']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ['凭证要素完整性', '100%完整', '96.7%', '否'],
        ['凭证编号连续性', '无断号重号', '正常', '是'],
        ['附件齐全率', '100%齐全', '95.2%', '否'],
        ['签字完备率', '100%完备', '98.5%', '否'],
        ['综合得分', '-', '28/30分', '-'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1.2 勾稽关系校验
    heading = doc.add_heading('1.2 勾稽关系校验', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('校验资产负债表与利润表、银行账与对账单、往来账与客户供应商对账的勾稽关系。')

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['勾稽项目', '校验规则', '校验结果', '差异金额']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ['资产负债表-利润表', '未分配利润变动=净利润', '通过', '0元'],
        ['银行账与对账单', '余额相等', '通过', '0元'],
        ['现金账与实盘', '余额相等', '未通过', '200元'],
        ['往来账对账', '账面与确认相等', '通过', '0元'],
    ]
    result_colors = ['E0FFE0', 'E0FFE0', 'FFE0E0', 'E0FFE0']
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if col_idx == 2:
                set_cell_shading(cell, result_colors[row_idx])

    doc.add_paragraph()

    # 1.3 异常交易识别
    heading = doc.add_heading('1.3 异常交易识别', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('识别大额异常交易、频繁整数交易、关联交易异常、资金回流检测等问题。')

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    headers = ['交易编号', '异常类型', '交易金额', '发生日期', '风险描述']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 5):
        for j, text in enumerate(['TXN001', '大额异常', '500,000', 'XXXX-XX-XX', '单笔超过月均50%']):
            cell = table.rows[i].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============= 第二章 =============
    heading = doc.add_heading('第二章  问题汇总', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 2.1 问题分类统计
    heading = doc.add_heading('2.1 问题分类统计', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['问题类别', '问题数量', '占比']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ['凭证规范性', '15', '45%'],
        ['勾稽关系', '8', '24%'],
        ['异常交易', '10', '30%'],
        ['其他问题', '1', '3%'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 2.2 重点问题清单
    heading = doc.add_heading('2.2 重点问题清单', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    headers = ['序号', '客户名称', '问题描述', '严重程度', '凭证编号']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2F5496')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 5):
        for j, text in enumerate(['1', '客户A', '缺少银行回单附件', '严重', '记-XXX']):
            cell = table.rows[i].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============= 第三章 =============
    heading = doc.add_heading('第三章  改进建议', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('本章针对诊断发现的问题，提供逐客户的改进方案和建议。')

    heading = doc.add_heading('3.1 严重问题改进方案（需立即整改）', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 3):
        p = doc.add_paragraph()
        run = p.add_run(f'客户A - 账务整改方案')
        run.font.bold = True
        run.font.size = Pt(12)

        p = doc.add_paragraph()
        p.add_run('存在问题：')
        p.add_run('[详细描述存在的问题]')

        p = doc.add_paragraph()
        p.add_run('整改措施：')
        p.add_run('[具体的整改步骤和方法]')

        p = doc.add_paragraph()
        p.add_run('完成时限：')
        p.add_run('[X个工作日内]')

        p = doc.add_paragraph()
        p.add_run('预计效果：')
        p.add_run('[整改后预期达到的效果]')

        doc.add_paragraph()

    heading = doc.add_heading('3.2 一般问题改进方案（限期整改）', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 3):
        p = doc.add_paragraph()
        run = p.add_run(f'客户B - 账务改进建议')
        run.font.bold = True
        run.font.size = Pt(12)

        p = doc.add_paragraph()
        p.add_run('存在问题：')
        p.add_run('[详细描述存在的问题]')

        p = doc.add_paragraph()
        p.add_run('改进建议：')
        p.add_run('[具体的改进建议]')

        doc.add_paragraph()

    # 保存
    output_path = 'D:/project/skills/代理记账/risk-monitor-guard/assets/report-templates/account-health-report-template.docx'
    doc.save(output_path)
    print(f'已生成: {output_path}')


def create_policy_change_bulletin():
    """创建政策变动快讯模板"""
    doc = Document()

    # 设置文档默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ============= 封面/标题区 =============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('XXX政策变动快讯')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('公告标题：关于XXXXX的政策公告')
    run.font.size = Pt(14)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('发布日期：XXXX年XX月XX日')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # 分隔线
    p = doc.add_paragraph()
    p.add_run('─' * 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ============= 第一章 =============
    heading = doc.add_heading('第一章  政策概述', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 政策基本信息表格
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    data = [
        ['政策名称', 'XXXXXXXX公告'],
        ['发文机关', '国家税务总局/财政部/XXXX'],
        ['文号', '公告XXXX年第X号'],
        ['发布日期', 'XXXX年XX月XX日'],
        ['生效日期', 'XXXX年XX月XX日'],
        ['政策类别', '优惠政策/征收管理/申报要求/发票管理'],
    ]
    for row_idx, row_data in enumerate(data):
        cell0 = table.rows[row_idx].cells[0]
        cell0.text = row_data[0]
        cell0.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell0, 'F2F2F2')
        cell1 = table.rows[row_idx].cells[1]
        cell1.text = row_data[1]

    doc.add_page_break()

    # ============= 第二章 =============
    heading = doc.add_heading('第二章  政策内容摘要', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 主要变化点
    heading = doc.add_heading('2.1 主要变化点', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('本政策的主要变化包括以下几个方面：')

    # 变化点列表
    for i in range(1, 4):
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(f'变化点{i}：')
        p.add_run('[详细描述变化内容]')

    doc.add_paragraph()

    # 新旧政策对比
    heading = doc.add_heading('2.2 新旧政策对比', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['项目', '旧政策', '新政策', '变化幅度']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i in range(1, 5):
        for j, text in enumerate(['项目1', '内容A', '内容B', '+X%']):
            cell = table.rows[i].cells[j]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============= 第三章 =============
    heading = doc.add_heading('第三章  影响分析', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 适用客户类型
    heading = doc.add_heading('3.1 适用客户类型', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('本政策主要影响以下类型的客户：')

    for i in range(1, 4):
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(f'客户类型{i}：')
        p.add_run('[描述适用条件]')

    doc.add_paragraph()

    # 预计影响金额
    heading = doc.add_heading('3.2 预计影响金额', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['影响类型', '客户数量', '预估影响金额']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ['减税影响（正面）', 'XX户', 'XX万元'],
        ['增税影响（负面）', 'XX户', 'XX万元'],
        ['中性影响', 'XX户', '-'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 风险提示
    heading = doc.add_heading('3.3 风险提示', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.add_run('特别提醒：')
    p.add_run('[提示在政策执行过程中需要注意的风险点]')

    doc.add_page_break()

    # ============= 第四章 =============
    heading = doc.add_heading('第四章  操作建议', level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 申报调整
    heading = doc.add_heading('4.1 申报调整', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 4):
        p = doc.add_paragraph()
        p.style = 'List Number'
        p.add_run(f'申报调整项{i}：')
        p.add_run('[描述需要调整的申报内容及具体操作方法]')

    doc.add_paragraph()

    # 资料准备
    heading = doc.add_heading('4.2 资料准备', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 4):
        p = doc.add_paragraph()
        p.style = 'List Number'
        p.add_run(f'资料准备项{i}：')
        p.add_run('[描述需要准备或补充的资料清单]')

    doc.add_paragraph()

    # 时间节点
    heading = doc.add_heading('4.3 时间节点', level=2)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['节点', '截止日期', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ['政策生效', 'XXXX-XX-XX', '新政策正式实施'],
        ['过渡期截止', 'XXXX-XX-XX', '过渡期结束'],
        ['首个执行节点', 'XXXX-XX-XX', '首个申报周期执行'],
        ['资料提交截止', 'XXXX-XX-XX', '需提交的备案资料截止日'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 页脚
    p = doc.add_paragraph()
    p.add_run('─' * 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('【风险提示】本快讯仅供参考，具体执行请以官方文件为准。如有疑问，请及时联系您的专属客服。')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'编制日期：XXXX年XX月XX日  |  编制部门：政策研究部  |  有效期至：XXXX年XX月XX日')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 保存
    output_path = 'D:/project/skills/代理记账/risk-monitor-guard/assets/report-templates/policy-change-bulletin-template.docx'
    doc.save(output_path)
    print(f'已生成: {output_path}')


if __name__ == "__main__":
    # 确保目录存在
    os.makedirs('D:/project/skills/代理记账/risk-monitor-guard/assets/report-templates', exist_ok=True)

    # 生成三个模板
    create_tax_risk_alert_report()
    create_account_health_report()
    create_policy_change_bulletin()

    print('\n所有报告模板已生成完成！')
