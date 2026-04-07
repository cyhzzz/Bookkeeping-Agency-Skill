"""
月度经营分析报告模板生成脚本
使用python-docx创建专业的DOCX格式报告模板

模板结构：
- 封面
- 目录
- 第一章 经营概况
- 第二章 税负分析
- 第三章 风险提示
- 第四章 下月注意事项
- 第五章 个性化财税建议
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_cover_page(doc):
    """创建封面"""
    # 添加空行
    for _ in range(6):
        doc.add_paragraph()

    # 公司名称占位符
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("[公司名称]")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    # 报告标题
    report_title = doc.add_paragraph()
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = report_title.add_run("X年X月经营分析报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    # 添加空行
    for _ in range(4):
        doc.add_paragraph()

    # 代账公司logo占位符
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_para.add_run("[代账公司Logo]")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 添加空行
    for _ in range(2):
        doc.add_paragraph()

    # 报告日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(14)

    # 添加空行
    for _ in range(2):
        doc.add_paragraph()

    # 保密声明
    confidentiality = doc.add_paragraph()
    confidentiality.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = confidentiality.add_run("【保密声明】")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    statement = doc.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = statement.add_run("本报告仅限收件人使用，未经授权不得转发或披露")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 分页
    doc.add_page_break()


def create_toc_page(doc):
    """创建目录页"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("目  录")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    toc_items = [
        ("第一章", "经营概况", "3"),
        ("第二章", "税负分析", "5"),
        ("第三章", "风险提示", "7"),
        ("第四章", "下月注意事项", "8"),
        ("第五章", "个性化财税建议", "9"),
    ]

    for chapter, content, page in toc_items:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(f"{chapter}  {content}")
        run.font.size = Pt(14)
        # 添加tab使页码在右侧
        para.paragraph_format.line_spacing = 1.5
        # 添加页码（简化的目录项）
        tab_run = para.add_run(f"  ..................  {page}")
        tab_run.font.size = Pt(14)

    doc.add_page_break()


def create_chapter1(doc):
    """创建第一章：经营概况"""
    # 章节标题
    title = doc.add_paragraph()
    run = title.add_run("第一章  经营概况")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(12)

    # 1.1 主要财务指标
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("1.1  主要财务指标")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    # 财务指标表格
    table = doc.add_table(rows=8, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ["指标", "本月", "本年累计", "环比", "同比"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 数据行
    data_rows = [
        ["营业收入", "¥{amount}", "¥{amount}", "{pct}", "{pct}"],
        ["营业成本", "¥{amount}", "¥{amount}", "{pct}", "{pct}"],
        ["毛利润", "¥{amount}", "¥{amount}", "{pct}", "{pct}"],
        ["毛利率", "{pct}", "{pct}", "-", "-"],
        ["期间费用", "¥{amount}", "¥{amount}", "{pct}", "{pct}"],
        ["净利润", "¥{amount}", "¥{amount}", "{pct}", "{pct}"],
        ["税负率", "{pct}", "-", "-", "-"],
    ]

    for row_idx, row_data in enumerate(data_rows, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()

    # 1.2 收入情况
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("1.2  收入情况")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("本月收入金额：")
    run.font.size = Pt(11)
    run = para.add_run("[请填写]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    para = doc.add_paragraph()
    run = para.add_run("环比变化：")
    run.font.size = Pt(11)
    run = para.add_run("[请填写]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    para = doc.add_paragraph()
    run = para.add_run("收入结构图描述：")
    run.font.size = Pt(11)
    run = para.add_run("[请在此描述收入构成图表]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 1.3 成本情况
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("1.3  成本情况")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[请填写本月成本情况分析]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 1.4 费用情况
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("1.4  费用情况")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[请填写本月费用情况分析]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 1.5 利润情况
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("1.5  利润情况")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[请填写本月利润情况分析]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_page_break()


def create_chapter2(doc):
    """创建第二章：税负分析"""
    title = doc.add_paragraph()
    run = title.add_run("第二章  税负分析")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(12)

    # 2.1 各税种税负情况
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("2.1  各税种税负情况")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    # 税负情况表格
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["税种", "本月税额", "税负率", "行业均值", "偏差"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    tax_types = [
        ["增值税", "¥{amount}", "{pct}%", "{pct}%", "{pct}%"],
        ["企业所得税", "¥{amount}", "{pct}%", "{pct}%", "{pct}%"],
        ["个人所得税", "¥{amount}", "{pct}%", "{pct}%", "{pct}%"],
        ["附加税费", "¥{amount}", "{pct}%", "{pct}%", "{pct}%"],
    ]

    for row_idx, row_data in enumerate(tax_types, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()

    # 2.2 增值税分析
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("2.2  增值税分析")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("本月应纳税额：")
    run.font.size = Pt(11)
    run = para.add_run("[请填写]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    para = doc.add_paragraph()
    run = para.add_run("税负率：")
    run.font.size = Pt(11)
    run = para.add_run("[请填写]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    para = doc.add_paragraph()
    run = para.add_run("行业对比：")
    run.font.size = Pt(11)
    run = para.add_run("[请填写]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 2.3 企业所得税分析
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("2.3  企业所得税分析")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[请填写企业所得税分析内容]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 2.4 综合税负率
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("2.4  综合税负率")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[请填写综合税负率分析内容]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 2.5 税负变化趋势
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("2.5  税负变化趋势")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("近6个月税负率折线图描述：")
    run.font.size = Pt(11)
    run = para.add_run("[请在此描述税负率变化趋势图表]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_page_break()


def create_chapter3(doc):
    """创建第三章：风险提示"""
    title = doc.add_paragraph()
    run = title.add_run("第三章  风险提示")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(12)

    # 风险等级说明
    note = doc.add_paragraph()
    run = note.add_run("风险等级说明：")
    run.font.size = Pt(11)
    run.font.bold = True
    run = note.add_run("高风险  ")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)
    run = note.add_run("中风险  ")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 153, 0)
    run = note.add_run("低风险")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 176, 80)

    # 3.1 票据异常汇总
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("3.1  票据异常汇总")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    # 异常票据表格
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["异常类型", "数量", "涉及金额", "风险等级"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    anomaly_data = [
        ["疑似假票", "0", "¥0", "-"],
        ["逾期票据", "0", "¥0", "-"],
    ]

    for row_idx, row_data in enumerate(anomaly_data, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()

    # 3.2 申报异常汇总
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("3.2  申报异常汇总")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[如有申报异常，请在此说明]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 3.3 政策变动影响
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("3.3  政策变动影响")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[如有政策变动影响，请在此说明]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 3.4 其他风险提示
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("3.4  其他风险提示")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[如有其他风险提示，请在此说明]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_page_break()


def create_chapter4(doc):
    """创建第四章：下月注意事项"""
    title = doc.add_paragraph()
    run = title.add_run("第四章  下月注意事项")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(12)

    # 4.1 即将到期事项
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("4.1  即将到期事项")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    # 申报日历表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["税种", "申报截止日期", "注意事项"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    filing_data = [
        ["增值税", "{date}", "{remark}"],
        ["企业所得税", "{date}", "{remark}"],
    ]

    for row_idx, row_data in enumerate(filing_data, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()

    # 4.2 需配合事项
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("4.2  需配合事项")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("□  ")
    run.font.size = Pt(11)
    run = para.add_run("[请填写需要客户配合的事项]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 4.3 政策动态
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("4.3  政策动态")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[请填写近期政策动态及影响]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_page_break()


def create_chapter5(doc):
    """创建第五章：个性化财税建议"""
    title = doc.add_paragraph()
    run = title.add_run("第五章  个性化财税建议")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(12)

    # 5.1 成本优化建议
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("5.1  成本优化建议")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[请根据客户情况填写成本优化建议]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 5.2 税务筹划建议
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("5.2  税务筹划建议")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[请根据客户情况填写税务筹划建议]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # 5.3 管理提升建议
    sub_title = doc.add_paragraph()
    run = sub_title.add_run("5.3  管理提升建议")
    run.font.size = Pt(14)
    run.font.bold = True
    sub_title.paragraph_format.space_before = Pt(6)

    para = doc.add_paragraph()
    run = para.add_run("[请根据客户情况填写管理提升建议]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_paragraph()
    doc.add_paragraph()

    # 页脚信息
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—" * 40)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    footer_info = doc.add_paragraph()
    footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_info.add_run(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    footer_author = doc.add_paragraph()
    footer_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_author.add_run("编制人：代账会计智能助手")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)


def create_report_template():
    """创建完整的月度经营分析报告模板"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # 创建各个章节
    create_cover_page(doc)
    create_toc_page(doc)
    create_chapter1(doc)
    create_chapter2(doc)
    create_chapter3(doc)
    create_chapter4(doc)
    create_chapter5(doc)

    return doc


def main():
    """主函数"""
    output_path = "D:/project/skills/代理记账/monthly-tax-helper/assets/report-templates/monthly-analysis-report-template.docx"

    print(f"正在生成月度经营分析报告模板...")
    doc = create_report_template()
    doc.save(output_path)
    print(f"模板已生成：{output_path}")


if __name__ == "__main__":
    main()
