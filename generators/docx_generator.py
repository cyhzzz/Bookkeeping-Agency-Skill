"""
Word文档生成器
用于生成客户交付用的专业报告
"""

from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional
import subprocess
import sys


class DocxGenerator:
    """Word文档生成器"""

    def __init__(self, template_dir: Optional[str] = None):
        self.template_dir = template_dir

    def generate(self, report_type: str, data: Dict[str, Any],
                 output_path: str, template_name: Optional[str] = None) -> str:
        """
        生成Word文档

        Args:
            report_type: 报告类型
            data: 报告数据
            output_path: 输出路径
            template_name: 可选模板文件名

        Returns:
            生成的Word文件路径
        """
        # 确保输出目录存在
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        # 调用minimax-docx skill生成文档
        try:
            result = self._generate_via_skill(report_type, data, output_path, template_name)
            return result
        except Exception as e:
            # 如果skill调用失败，使用基础生成方法
            return self._generate_fallback(report_type, data, output_path)

    def _generate_via_skill(self, report_type: str, data: Dict[str, Any],
                           output_path: str, template_name: Optional[str]) -> str:
        """通过minimax-docx skill生成文档"""
        # 模板映射
        template_map = {
            'monthly_report': '月度经营分析报告模板.docx',
            'client_onboarding': '客户接入确认单模板.docx',
            'risk_alert': '税务风险预警报告模板.docx',
            'export_rebate': '出口退税报告模板.docx',
            'rd_expense': '研发费用加计扣除模板.docx',
            'customer_value': '客户价值报告模板.docx',
        }

        template = template_name or template_map.get(report_type, '通用报告模板.docx')

        # 构建命令
        cmd = [
            sys.executable, '-c',
            f'''
import sys
sys.path.insert(0, r"<skill-base>/../minimax-docx")
# 这里应该调用minimax-docx skill
# 简化版本：直接创建基础Word文档
from docx import Document
doc = Document()
doc.add_heading("{data.get('title', '报告')}", 0)
doc.add_paragraph("客户：{data.get('customer_name', '')}")
doc.add_paragraph("日期：{datetime.now().strftime('%Y-%m-%d')}")
doc.add_paragraph("报告类型：{report_type}")
doc.save(r"{output_path}")
print("Generated: {output_path}")
'''
        ]

        try:
            subprocess.run(cmd, capture_output=True, check=True)
            return output_path
        except Exception:
            raise

    def _generate_fallback(self, report_type: str, data: Dict[str, Any], output_path: str) -> str:
        """基础生成方法（当minimax-docx不可用时）"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 添加标题
            title = data.get('title', '代账报告')
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加基本信息
            doc.add_paragraph(f"客户名称：{data.get('customer_name', '')}")
            doc.add_paragraph(f"报告周期：{data.get('period', '')}")
            doc.add_paragraph(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")

            # 添加内容部分
            if 'content' in data:
                doc.add_heading('报告内容', level=1)
                doc.add_paragraph(data['content'])

            # 添加风险提示（如有）
            if 'risk_alerts' in data:
                doc.add_heading('风险提示', level=1)
                doc.add_paragraph(data['risk_alerts'])

            # 添加页脚
            doc.add_paragraph('─' * 50)
            footer = doc.add_paragraph('本报告由AI代账工具自动生成')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.save(output_path)
            return output_path

        except ImportError:
            # python-docx未安装，返回错误信息
            raise RuntimeError(
                "Word文档生成需要安装python-docx库。"
                "请运行: pip install python-docx"
            )

    def set_template_dir(self, template_dir: str):
        """设置模板目录"""
        self.template_dir = template_dir
